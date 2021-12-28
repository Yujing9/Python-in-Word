#Reading Word Documents
import docx
doc = docx.Document('demo.docx')
len(doc.paragraphs)
doc.paragraphs[0].text
doc.paragraphs[1].text
len(doc.paragraphs[1].runs)
doc.paragraphs[1].runs[0].text
doc.paragraphs[1].runs[1].text
doc.paragraphs[1].runs[2].text
doc.paragraphs[1].runs[3].text
#Getting the Full Text from a .docx File
import docx
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        #fullText.append(' ' + para.text)
    return '\n'.join(fullText)
    #return '\n\n'.join(fullText)
print(getText('demo.docx'))
#Styling Paragraph and Run Objects
#Run Attributes
doc = docx.Document('demo.docx')
doc.paragraphs[0].text
doc.paragraphs[0].style
doc.paragraphs[0].style = 'Normal'
doc.paragraphs[1].text
#一个paragraph有多个run，每个run有对应的相应的attribute
(doc.paragraphs[1].runs[0].text, doc.paragraphs[1].runs[1].text, doc.
paragraphs[1].runs[2].text, doc.paragraphs[1].runs[3].text)
('A plain paragraph with some ', 'bold', ' and some ', 'italic')
doc.paragraphs[1].runs[0].style = 'QuoteChar'
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True
doc.save('restyled.docx')
#Writing Word Documents
#automate to create a new file
import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')
doc.save('helloworld.docx')
#--------------------------------------------------------------------
import docx
doc = docx.Document()
doc.add_paragraph('Hello world!')
paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is a yet another paragraph.')
paraObj1.add_run(' This text is being added to the second paragraph.')
doc.save('multipleParagraphs.docx')
#Adding Headings
doc = docx.Document()
doc.add_heading('Header 0', 0)
doc.add_heading('Header 1', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('Header 3', 3)
doc.add_heading('Header 4', 4)
doc.save('headings.docx')
#Adding Line and Page Breaks
doc = docx.Document()
doc.add_paragraph('This is on the first page!')
#Adding line
doc.paragraphs[0].runs[0].add_break()
#Adding page
doc.paragraphs[0].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
doc.add_paragraph('This is on the second page!')
doc.save('twoPage.docx')
#Adding Pictures   
import docx
doc = docx.Document()
doc.add_picture('zophie.png', width=docx.shared.Inches(1),height=docx.shared.Cm(4))
doc.save('picture.docx')

